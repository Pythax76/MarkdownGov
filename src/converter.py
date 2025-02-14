import logging
import re
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE   
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import yaml
from datetime import datetime


# Configure Advanced Logging
log_file = "conversion_log.txt"
logging.basicConfig(filename=log_file, level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")

class MarkdownToWordConverter:
    def __init__(self):
        self.detected_title = None  # Store the detected title

    def convert(self, template_path, markdown_path, output_dir):
        """Converts a Markdown file into a Word document using a template."""
        doc = Document(template_path)

        # Generate output filename
        base_name = os.path.splitext(os.path.basename(markdown_path))[0]
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        output_filename = f"{timestamp}_{base_name}.docx"
        output_path = os.path.join(output_dir, output_filename)

        logging.debug(f"Generated output file: {output_filename}")

        # Read Markdown content
        with open(markdown_path, 'r', encoding='utf-8') as md_file:
            md_content = md_file.readlines()

        logging.debug("Starting Markdown to Word conversion.")

        # ✅ Extract metadata before processing the document
        metadata = self._extract_metadata(md_content)

        # ✅ Ensure title is correctly assigned (Use metadata or fallback)
        if metadata["Title"] == "-unassigned-":
            metadata, clean_md_content = self._detect_title(md_content, metadata)
        else:
            clean_md_content = md_content  # Use original content if title was from metadata

        title_text = metadata["Title"]

        # ✅ Apply title to the document using "Title" style
        doc.add_paragraph(title_text, style="Title")
        logging.debug(f"Applied 'Title' style to: {title_text}")

        # ✅ Assign the title to the Word document's metadata
        doc.core_properties.title = title_text
        logging.debug(f"Assigned document metadata title: {title_text}")

        # ✅ Convert Markdown to Word format
        self._parse_markdown_to_word(clean_md_content, doc)

        # ✅ Apply metadata to Word document
        self._apply_metadata(doc, metadata)

        # ✅ Save final Word document
        doc.save(output_path)
        logging.debug(f"Word file saved at: {output_path}")

        return output_path


    def _detect_title(self, md_lines, metadata):
        """Detects the title from the Markdown file and removes `===`."""
        clean_lines = []

        for i in range(len(md_lines) - 1):
            line = md_lines[i].strip()
            next_line = md_lines[i + 1].strip()

            # ✅ Detect a title only when underlined with "==="
            if next_line.startswith("=") and len(next_line) >= 3:
                logging.debug(f"Title detected: {line}")
                metadata["Title"] = line  # ✅ Assign title directly to metadata
                return metadata, md_lines[i + 2 :]  # ✅ Stop iteration & return remaining text

            clean_lines.append(line)

        return metadata, clean_lines  # ✅ If no title detected, return default metadata

    def _parse_markdown_to_word(self, md_lines, doc):
        """Parses Markdown content, applies formatting, and sets indentation."""
        current_indent = 0.0  # Default indentation (Heading 1 starts at 0.0")

        # Ensure "Code" Character Style Exists
        styles = doc.styles
        if "Inline Code" not in styles:
            code_style = styles.add_style("Inline Code", WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = "Courier New"
            code_style.font.size = Inches(0.1)  # Adjust size if needed

        for line in md_lines:
            line = line.strip()

            if not line:
                continue  # Skip empty lines

            # Detect headings and apply correct indentation
            match = re.match(r'^(#{1,6})\s*(.*)', line)
            if match:
                level = len(match.group(1))  # Count number of '#' to determine heading level
                text = match.group(2)

                # Define indentation based on heading level
                heading_indent = (level - 1) * 0.25  # 0.25" per level
                current_indent = heading_indent  # Store for body text under this heading

                heading = doc.add_paragraph(text, style=f"Heading {level}")
                heading.paragraph_format.left_indent = Inches(heading_indent)

                logging.debug(f"Assigned Heading {level} with {heading_indent}\" indent: {text}")
                continue

            # Apply the current left indentation to all text under the heading
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(current_indent)

            # Fix Inline Code (Use CHARACTER Style Instead of PARAGRAPH Style)
            inline_code_matches = re.findall(r'`(.*?)`', line)
            if inline_code_matches:
                for code_snippet in inline_code_matches:
                    run = paragraph.add_run(code_snippet)
                    run.style = "Inline Code"  # Use a character style, not paragraph
                    logging.debug(f"Applied Inline Code: {code_snippet}")

                line = re.sub(r'`(.*?)`', '', line)  # Remove inline code from the sentence

            # Handle Bold, Italic, and Bold+Italic
            line = re.sub(r'\*\*\*(.*?)\*\*\*', r'<b><i>\1</i></b>', line)  # Bold+Italic
            line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)  # Bold
            line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)  # Italic

            tokens = re.split(r'(<b>|</b>|<i>|</i>)', line)
            is_bold, is_italic = False, False

            for token in tokens:
                if token == "<b>":
                    is_bold = True
                elif token == "</b>":
                    is_bold = False
                elif token == "<i>":
                    is_italic = True
                elif token == "</i>":
                    is_italic = False
                else:
                    run = paragraph.add_run(token)
                    run.bold = is_bold
                    run.italic = is_italic

            logging.debug(f"Formatted Inline Text: {line}")
            
    def _extract_metadata(self, md_lines):
        """Extracts metadata from YAML-like front matter in the Markdown file."""
        metadata = {
            "Title": "-unassigned-",
            "Document ID": "-unassigned-",
            "Facility": "-unassigned-",
            "Version": "-unassigned-",
            "Category": "-unassigned-",
            "Content": "-unassigned-",
            "Author": "-unassigned-"
        }

        yaml_content = []
        in_yaml_block = False

        for line in md_lines:
            line = line.strip()

            # Detect start of YAML front matter
            if line == "---":
                in_yaml_block = not in_yaml_block  # Toggle state
                continue

            if in_yaml_block:
                yaml_content.append(line)

        # Parse YAML metadata if found
        if yaml_content:
            try:
                extracted_metadata = yaml.safe_load("\n".join(yaml_content))
                for key in extracted_metadata:
                    metadata[key] = extracted_metadata[key]  # Override defaults if a value is present
                logging.debug(f"Extracted Metadata: {metadata}")
            except yaml.YAMLError as e:
                logging.warning(f"Failed to parse metadata: {e}")

        return metadata


    def _apply_metadata(self, doc, metadata):
        """Applies document metadata, preserving '-unassigned-' when no value exists."""
        properties = doc.core_properties

        # Assign metadata fields, keeping '-unassigned-' as a default
        properties.title = metadata.get("Title", "-unassigned-")
        properties.author = metadata.get("Author", "-unassigned-")
        properties.subject = metadata.get("Category", "-unassigned-")
        properties.version = metadata.get("Version", "-unassigned-")

        # Ensure custom fields like Document ID, Facility, etc., remain "-unassigned-" if no value is provided
        doc.add_paragraph(f"Document ID: {metadata.get('Document ID', '-unassigned-')}", style="Normal")
        doc.add_paragraph(f"Facility: {metadata.get('Facility', '-unassigned-')}", style="Normal")
        doc.add_paragraph(f"Content Category: {metadata.get('Content', '-unassigned-')}", style="Normal")

        logging.debug(f"Metadata applied: Title = {properties.title}, Author = {properties.author}, Version = {properties.version}")
